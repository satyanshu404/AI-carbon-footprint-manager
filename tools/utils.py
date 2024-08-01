import sys
import os

# Add the root directory to sys.path
sys.path.append(os.path.dirname(os.path.dirname(os.path.abspath(__file__))))

import re
import ast
import pandas as pd
import PyPDF2
import docx
import json
from bs4 import BeautifulSoup
from openai import OpenAI
import base64
import mimetypes
import requests
import constants
import importlib.util
from openpyxl import load_workbook
from typing import Optional
from groq import Groq
from prompts import prompts
import tools.reteriver as rt
from langchain.text_splitter import MarkdownTextSplitter
from llama_index.core import VectorStoreIndex, get_response_synthesizer
from llama_index.core import Document
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.retrievers import VectorIndexRetriever
from llama_index.core.query_engine import RetrieverQueryEngine
from llama_index.core.postprocessor import SimilarityPostprocessor

from dotenv import load_dotenv


load_dotenv(override=True)

class ReadFiles:
    '''Reads the content of a file based on its extension and returns the content as a string or a dataframe.'''
    def __init__(self):
        pass

    def get_extension(self, file_path:str) -> str:
        '''Returns the extension of a file.'''
        return os.path.splitext(file_path)[1].lower()
    
    def read_file(self, file_path:str) -> str:
        '''Reads the content of a file based on its extension and returns the content as a string or a dataframe.'''
        print(self.get_extension(file_path))
        if self.get_extension(file_path) == '.txt':
            return self.read_txt(file_path)
        elif self.get_extension(file_path) == '.pdf':
            return self.read_pdf(file_path)
        elif self.get_extension(file_path) == '.docx':
            return self.read_docx(file_path)
        elif self.get_extension(file_path) in ['.csv', '.xlsx', '.xls']: 
            return self.read_excel_or_csv(file_path)
        elif self.get_extension(file_path) == '.html':
            return self.read_html(file_path)
        elif self.get_extension(file_path) == '.json':
            json_data = self.read_json(file_path)
            return json.dumps(json_data, indent=4)
        else:
            raise ValueError(f"Unsupported file type: {self.get_extension(file_path)}")
    
    def read_txt(self, file_path:str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            return file.read()
    
    def read_pdf(self, file_path: str) -> str:
        with open(file_path, 'rb') as file:
            reader = PyPDF2.PdfReader(file)
            text = ''
            for page in reader.pages:
                text += page.extract_text()
        return text
    
    def read_docx(self, file_path:str) -> str:
        doc = docx.Document(file_path)
        return '\n'.join([paragraph.text for paragraph in doc.paragraphs])
    
    def read_xlsx(self, file_path:str) -> str:
        '''Reads the content of an excel file and returns the content as a formated string.
        Example:
            1 2 3 A B
            ---------
            4 5 6 C D
            ---------
        '''
        workbook = load_workbook(filename=file_path)
        sheet = workbook.active
        
        data:list = []
        for row in sheet.iter_rows(values_only=True):
            if any(cell is not None and cell != "" for cell in row):
                data.append(row)

        string: str = ""
        for row in data:
            row_data =  str(row)
            row_data = row_data.replace("None", "")
            row_data = row_data.replace("(", "").replace(")", "")
            row_data = row_data.replace(",", " ")
            string += row_data + f"\n{50*'-'}" + "\n"
            
        return string
    
    def read_excel_or_csv(self, file_path:str) -> str:
        flag = False
        if self.get_extension(file_path) == '.csv':
            df = pd.read_csv(file_path)
            path = file_path.split('.')[0] + '.xlsx'
            df.to_excel(path, index=False)
            file_path = path
            flag = True
            
        data: str = self.read_xlsx(file_path)
        if flag:    
            os.remove(file_path)
        return data
    
    def read_json(self, file_path:str) -> str:
        with open(file_path, 'r') as file:
            return json.load(file)
    
    def read_html(self, file_path:str) -> str:
        with open(file_path, 'r', encoding='utf-8') as file:
            soup = BeautifulSoup(file, 'html.parser')
            return soup.get_text()
        
class ReadImage:
    def __init__(self, model_name:str = constants.Constants.MODEL_NAME):
        self.model_name:str = model_name
        self.client = OpenAI()
    
    def image_to_base64(self, image_path:str):
        '''Converts an image to base64 string'''
        mime_type, _ = mimetypes.guess_type(image_path)
        if not mime_type or not mime_type.startswith('image'):
            raise ValueError("The file type is not recognized as an image")
        
        with open(image_path, 'rb') as image_file:
            encoded_string = base64.b64encode(image_file.read()).decode('utf-8')
        
        image_base64 = f"data:{mime_type};base64,{encoded_string}"
        return image_base64

    def read_image(self, image_path:str) -> str:
        '''Reads an image using OPENAI'''
        try:
            base64_string = self.image_to_base64(image_path)
            messages=[
                    {"role": "system", "content": "You are a helpful assistant that responds in Markdown."},
                    {"role": "user", "content": [
                        {"type": "text", "text": "Carefully analyze the image and extract its content, then present the extracted information in JSON format., nothing else"},
                        {"type": "image_url", 
                         "image_url": {
                            "url": base64_string}
                        }
                ]}]
            
            response = self.client.chat.completions.create(
                model=self.model_name,
                messages=messages
            )
            return response.choices[0].message.content
        except Exception as e:
            return f"A error occurred while reading the image: {str(e)}"

class GoogleSearch:
    def __init__(self):
        self.key = os.getenv('GOOGLE_API_KEY')
        self.cx = os.getenv('GOOGLE_SEARCH_ENGINE_ID')
        self.url = constants.GoogleSearchConstants.GOOGLE_SEARCH_URL
        self.save_file_path = constants.GoogleSearchConstants.GOOGLE_SEARCH_RESULTS_PATH

    def build_payloads(self, 
                       query: str, 
                       start:int=1, 
                       num:int=10,
                       **params)->dict:
        ''' Build the payloads for the Google Search API'''

        payload = {
            'key': self.key,
            'cx': self.cx,
            'q': query,
            'start': start,
            'num': num,
        }
        payload.update(params)
        return payload

    def make_request(self, payload: dict) -> dict:
        '''Make a request to the Google Search API'''
        response = requests.get(self.url, params=payload)
        return response.json()

    def clean_file_name(self, file_name:str)->str:
        '''Clean the file name'''
        # remove special characters
        file_name = file_name.lower()
        return re.sub(r'[^\w\s-]', '', file_name).strip().replace(' ', '_')

    def search(self,
               query:str,
               restult_count:int=10)->str:
        '''Search Google and return status of the search'''
        items = []
        reminder = restult_count%10
        if reminder > 0:
            pages = restult_count//10 + 1
        else:
            pages = restult_count//10 

        try:
            for i in range(pages):
                if pages == i+1 and reminder > 0:
                    payload = self.build_payloads(query, start=(i+1)*10, num=reminder)
                else:
                    payload = self.build_payloads(query, start=(i+1)*10)
                # print(payload)
                response = self.make_request(payload)
                # print(response)
                items.extend(response['items'])

            query_string_cleaned = self.clean_file_name(query)
            df = pd.json_normalize(items)
            search_urls = [url for url in df['link']]
            # save_file_path = '{0}/Search_resutls_{1}.csv'.format(self.save_file_path, query_string_cleaned)
            # df.to_csv(save_file_path, index=False)
            return search_urls
        except Exception as e:
            return f'''Error occured during search: {e}'''

class GroqModel:
    '''Groq model class'''
    def __init__(self, model_name:str = constants.GroqModelConstants.MODEL_NAME):
        self.client = Groq()
        self.model = model_name
    def get_completion(self, messages, **kwargs) -> str:
        '''Get completion from the model'''
        chat_completion = self.client.chat.completions.create(
            messages=messages,
            model=self.model,
            **kwargs
        )
        if chat_completion:
            # print(chat_completion)
            return chat_completion.choices[0].message.content

class GptModel:
    '''Gpt model class'''
    def __init__(self, model_name:str = constants.Constants.MODEL_NAME):
        self.client = OpenAI()
        self.model = model_name
    
    def create_prompt(self, prompt:str) -> list[dict]:
        return [{'role': 'system', 'content': prompts.get_system_prompt_for_ai_assistant()},
                {'role': 'user', 'content': prompt}]
    
    def get_completion(self, prompt:str, **kwargs) -> str:
        '''Get completion from the model'''
        completion = self.client.chat.completions.create(
            model=self.model,
            messages= self.create_prompt(prompt),
            **kwargs
        )
        return completion.choices[0].message.content

class CodeExecuter:
    def __init__(self):
        pass
    
    def generate_code(self, prompt:str, filename:str, function_name:str):
        '''Generate code from the messages'''
        groq = GroqModel()
        messages = [
            {'role': 'system', 'content': 'You are a Python coder. Just give the python code in text, nothing else, no notations, not any asterisk, nothing. Just the code.'},
            {'role': 'system', 'content': f'The pyhton function must have the name {function_name}'},
            {'role': 'user', 'content': prompt}]
        code =  groq.get_completion(messages)

        with open(filename, 'w+') as file:
            file.write(code)

    def load_function_from_file(self, filename:str, function_name:str):
        '''Load the function from a file'''
        spec = importlib.util.spec_from_file_location(function_name, filename)
        module = importlib.util.module_from_spec(spec)
        spec.loader.exec_module(module)
        return getattr(module, function_name)
    
    def execute_code(self, prompt:str, function_name:str, **kwargs):
        '''Execute the generated function'''
        try:
            code_filename = f"{constants.CodeExecuterConstants.BASEFILEPATH}/generated_code.py"
            self.generate_code(prompt, code_filename, function_name)
            generated_function = self.load_function_from_file(code_filename, function_name)
            # result = generated_function(**kwargs)
            if kwargs:
                result = generated_function(**kwargs)
            else:
                result = generated_function()
            return result
        except Exception as e:
            return f"An error occurred while executing the code: {str(e)}"
        
class Retriever:
    def __init__(self):
        self.RF = ReadFiles()

    def create_documents(self, file_paths: list[str]) -> list[Document]:
        '''Create the documents from the file paths'''
        return [Document(text=self.RF.read_file(file_path)) for file_path in file_paths]
    
    def create_retriever(self, documents: list[Document]) -> VectorStoreIndex:
        '''Create the retriever'''
        index = VectorStoreIndex.from_documents(documents)
        retriever = VectorIndexRetriever(
                    index=index,
                    similarity_top_k=constants.RetrieverConstants.SIMILARITY_TOP_K,
                )
        return retriever
    
    def query_engine(self, query:str, retriever:VectorIndexRetriever) -> RetrieverQueryEngine:
        '''create the query engine and call the engine for the query'''
        response_synthesizer = get_response_synthesizer()
        query_engine = RetrieverQueryEngine(
            retriever=retriever,
            response_synthesizer=response_synthesizer,
            node_postprocessors=[SimilarityPostprocessor(similarity_cutoff=constants.RetrieverConstants.SIMILARITY_CUTOFF)],
        )
        response = query_engine.query(query)
        return response

    def retrieve(self, query:str, file_paths: list[str]) -> list[str] | str:
        '''Retrive the relevant chunks from the files at the based on the query and returns the chunks as a list of strings which are important to the query.'''
        try:
            print("Retrieving chunks...")
            documents = self.create_documents(file_paths)
            retriever = self.create_retriever(documents)
            response: RetrieverQueryEngine = self.query_engine(query, retriever)
            # response = self.query_engine(query, retriever)

            print(f"Got {len(response.source_nodes)} chunks.")
            return [doc.text for doc in response.source_nodes]
        except Exception as e:
            return f"Got an error while reteriving the chunks: {e}"

class WebScraper:
    def __init__(self):
        self.header = {'User-Agent': constants.WebScraperConstants.USERAGENT}

        if not os.path.exists(constants.WebScraperConstants.SCRAPER_FILE_LOCATION):
                os.makedirs(constants.WebScraperConstants.SCRAPER_FILE_LOCATION)

    def make_request(self, url:str):
        '''Make a request to the url'''
        response = requests.get(url, headers=self.header)
        return response
    
    def scrape(self, url:str, query:str = None) -> tuple[str, list[str]] | str:
        '''Scrape the content aof the url and return the relevant content based on the query'''
        try:
            response = self.make_request(url)

            if len(url.split('/')[-1]) == 0:                
                file_path = os.path.join(constants.WebScraperConstants.SCRAPER_FILE_LOCATION, url.split('/')[-2])
            else:
                file_path = os.path.join(constants.WebScraperConstants.SCRAPER_FILE_LOCATION, url.split('/')[-1])

            if not file_path.endswith(('.pdf','.xlsx','.xls','.csv','.docx','.txt')):
                file_path += '.html'

            with open(file_path, 'wb') as file:
                file.write(response.content)
            if query:
                return rt.RetrieverRouter().route(query, [file_path])
            
            return ReadFiles().read_file(file_path)
        except Exception as e:
            return f"An error occurred while scraping the content: {str(e)}"

class JsonFormatter:
    '''Format the json file for the langchain agent
    Description:
        - all the curly braces {} are replaced with dict()
        - all the square brackets [] are replaced with list()
        Note: Having {} in json file cause conflict with the langchain agent's input
    '''
    def __init__(self):
        self.save_dir = constants.FormatJsonConstants.SAVE_DIRECTORY

    def to_string(self, json_string: str) -> str:
        '''load json as json string'''
        return json.dumps(json_string, indent=4)
    
    def format_text(self, text:str) -> str:
        '''Format the text'''
        text = text.replace('{', 'dict(')
        text = text.replace('}', ')')
        text = text.replace('[', 'list(')
        text = text.replace(']', ')')
        return text
    
    def save_file(self, text:str, file_name:str) -> str:
        '''Save the file'''
        file_path = self.save_dir.format(file_name)
        with open(file_path, 'w') as file:
            file.write(text)

    def format_json(self, file_path:str, file_name: str) -> str:
        '''Format the json file for the langchain agent
        Description:
            - all the curly braces {} are replaced with 'dict()'
            - all the square brackets [] are replaced with 'list()'
        Note: Having {} in json file cause conflict with the langchain agent's input
        '''
        try:
            json_string = json.load(open(file_path))
            text = self.to_string(json_string)
            formated_text = self.format_text(text)

            self.save_file(formated_text, file_name)
            return f"Json file saved successfully at {self.save_dir.format(file_name)}"
        except Exception as e:
            return f"An error occurred while formatting the json file: {str(e)}"
        
class TextSplitter:
    '''Split the text into chunks using the markdown text splitter'''
    def __init__(self,
                 chunk_size:int = constants.TextSplitterConstants.CHUNK_SIZE,
                 chunk_overlap:int = constants.TextSplitterConstants.CHUNK_OVERLAP):
        self.chunk_size = chunk_size
        self.chunk_overlap = chunk_overlap
    
    def split_text(self, document: str) -> list[str]:
        '''Split the text into chunks using the markdown text splitter'''
        chunks = []
        splitter = MarkdownTextSplitter(chunk_size=self.chunk_size, chunk_overlap=self.chunk_overlap)
        chunks = splitter.split_text(document)
        return chunks
    
class FormatResponse:
    '''Format the response of the llm calls'''
    def __init__(self):
        pass
    
    def format_response(self, response:str, pattern:str) -> str:
        """
        Extract the json/python code from the response.
        Args:
        - response (str): Input text containing the response surrounded by backticks and pattern.
        - pattern (str): The pattern (e.g., "python") to match and remove along with the backticks.
        """
        # Construct regex pattern to match backticks and the specified pattern
        pattern_regex = rf'^```(?:{pattern})?\n|\n```$'
        
        # Remove surrounding ``` and the specified pattern
        cleaned_text = re.sub(pattern_regex, '', response, flags=re.MULTILINE)
        
       
        python_list = ast.literal_eval(cleaned_text.strip())
        return python_list
    
class LLMRetriever:
    def __init__(self):
        self.model = GptModel()
        self.RF = ReadFiles()
        self.text_splitter = TextSplitter(chunk_overlap=constants.LLMRetrieverConstants.CHUNK_OVERLAP,
                                          chunk_size=constants.LLMRetrieverConstants.CHUNK_SIZE)

    def make_prompt(self, query:str, context: str) -> list[dict]:
        '''Make the prompt for the LLM'''
        return prompts.get_system_prompt_for_llm_retriever().format(context, query)

    def retrieve(self, query:str, file_paths: list[str]) -> Optional[list[str]]:
        '''Retrieve the relevant chunks from the documents based on the query, context and instruction.'''
        documents:list = [self.RF.read_file(file_path) for file_path in file_paths]
        
        data = []
        for document in documents:
            chunks: list[str] = self.text_splitter.split_text(document)
            print(f"Got {len(chunks)} chunks.")
            messages: list[str] = [self.make_prompt(query, chunk) for chunk in chunks]  
            responses: list[str] = [self.model.get_completion(message) for message in messages]
            data.append(responses)

        return data

class JsonWriter:
    '''Write the json file'''
    def __init__(self):
        pass

    def read_json_file(self, file_path: str) -> list[dict]:
        '''Read the json file'''
        if os.path.exists(file_path):
            with open(file_path, 'r') as file:
                data = json.load(file)
        else:
            data = []
        return data
    
    def write_json_file(self, file_path:str, data: list[dict]) -> None:
        '''Write the json file'''
        with open(file_path, 'w') as file:
            json.dump(data, file, indent=4)
        
    def write(self, data: str, file_path:str) -> str:
        '''Write the json file'''
        try:
            existing_data = self.read_json_file(file_path)
            existing_data.extend(json.loads(data))   
            self.write_json_file(file_path, existing_data)
        except Exception as e:
            return f"An error occurred while writing the data: {str(e)}"
